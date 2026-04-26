"""
Parser D — EVEBI Nachhaltigkeit.docx (Ollama-unterstützt, Confidence < 1.0)

Format: DOCX (Word-Dokument), Export "Ökobilanzierung / Nachhaltigkeit"
Erkennung: Dateiname enthält "nachhalt" oder Tabelle mit "PLAN / PLUS / PREMIUM"

Strategie:
  1. python-docx liest alle Tabellen
  2. Heuristischer Matcher sucht die QNG-Ergebnistabelle
     (Spaltenköpfe: PLAN, PLUS, PREMIUM, Einheit)
  3. Falls Heuristik scheitert → Ollama-Fallback mit dem gesamten Text

Confidence-Zuordnung:
  - Heuristischer Match: 0.92 (manuelle Auditor-Prüfung empfohlen)
  - Ollama-Extraktion:   0.75 (Auditor-Prüfung Pflicht)
  - Keine Extraktion:    0.0  (Datei nicht erkannt)
"""

import json
import re
import urllib.request
import urllib.error
from typing import Any

try:
    from docx import Document  # type: ignore
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


# Tabellenzeilen die wir suchen (case-insensitive, Teilstring-Match)
_GWP_LABELS  = ["treibhausgasemissionen", "gwp", "co₂", "co2"]
_PENE_LABELS = ["primärenergiebedarf", "primärenergie", "pene", "nicht erneuerbar"]
_BGF_LABELS  = ["bgf", "bruttogeschoss", "gross floor"]
_NRF_LABELS  = ["nrf", "nettoraum", "net floor", "nrf(r)"]

OLLAMA_URL  = "http://localhost:11434/api/generate"
OLLAMA_MODEL = "llama3"

_UNIT_NORM = {
    "kg co₂-äqu./(m²nrfa)": "kg_co2_m2a",
    "kwh/(m²a)":             "kwh_m2a",
    "m²":                    "m2",
}


def _row_matches(row_text: str, labels: list[str]) -> bool:
    lower = row_text.lower()
    return any(lbl in lower for lbl in labels)


def _cell_float(text: str) -> float | None:
    """Bereinigt einen Zelltext und versucht Float-Konvertierung."""
    cleaned = re.sub(r"[^\d,.\-]", "", text.strip())
    cleaned = cleaned.replace(",", ".")
    try:
        return float(cleaned)
    except ValueError:
        return None


def _parse_qng_table(table: Any) -> dict[str, float | None]:
    """
    Versucht die QNG-Ergebnistabelle (PLAN/PLUS/PREMIUM/Einheit) zu parsen.
    Gibt ein Dict mit den gefundenen Werten zurück, oder {} wenn Tabelle nicht passt.
    """
    rows = table.rows
    if len(rows) < 2:
        return {}

    # Kopfzeile finden: enthält PLAN + PLUS + PREMIUM
    header_idx = None
    plan_col = plus_col = premium_col = -1

    for i, row in enumerate(rows):
        texts = [cell.text.strip().upper() for cell in row.cells]
        if "PLAN" in texts and "PLUS" in texts and "PREMIUM" in texts:
            header_idx = i
            plan_col    = texts.index("PLAN")
            plus_col    = texts.index("PLUS")
            premium_col = texts.index("PREMIUM")
            break

    if header_idx is None:
        return {}

    result: dict[str, float | None] = {}

    for row in rows[header_idx + 1:]:
        cells = [cell.text.strip() for cell in row.cells]
        if not cells or not cells[0]:
            continue

        label = cells[0]
        plan_val = _cell_float(cells[plan_col]) if plan_col < len(cells) else None

        if _row_matches(label, _GWP_LABELS):
            result["gwp_gesamt"] = plan_val
            if plus_col < len(cells):
                result["qng_plus_gwp"] = _cell_float(cells[plus_col])
            if premium_col < len(cells):
                result["qng_premium_gwp"] = _cell_float(cells[premium_col])

        elif _row_matches(label, _PENE_LABELS):
            result["pene_gesamt"] = plan_val

        elif _row_matches(label, _BGF_LABELS):
            result["bgf_m2"] = plan_val

        elif _row_matches(label, _NRF_LABELS):
            result["nrf_m2"] = plan_val

    return result


def _extract_full_text(doc: Any) -> str:
    """Gibt den gesamten lesbaren Textinhalt des DOCX zurück."""
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def _ollama_extract(text: str) -> tuple[dict[str, Any], float]:
    """
    Sendet den Dokumenttext an Ollama und bittet um JSON-Extraktion.
    Gibt (ki_extrahiert, confidence) zurück.
    confidence = 0.75 wenn Ollama antwortet, sonst 0.0.
    """
    prompt = f"""Du analysierst einen EVEBI-Nachhaltigkeitsbericht.
Extrahiere folgende Werte als JSON-Objekt (nur die Zahlen, keine Einheiten):
- gwp_gesamt: GWP / Treibhausgasemissionen PLAN [kg CO2-Äqu./(m²NRFa)]
- pene_gesamt: Primärenergiebedarf nicht erneuerbar PLAN [kWh/(m²a)]
- bgf_m2: BGF nach DIN 277 [m²]
- nrf_m2: NRF nach DIN 277 [m²]
- qng_plus_gwp: GWP PLUS-Anforderung [kg CO2-Äqu./(m²NRFa)] (falls vorhanden)
- qng_premium_gwp: GWP PREMIUM-Anforderung (falls vorhanden)

Dokumenttext:
{text[:4000]}

Antwort NUR als JSON, kein Text davor oder danach:
"""
    payload = json.dumps({
        "model": OLLAMA_MODEL,
        "prompt": prompt,
        "stream": False,
    }).encode()

    try:
        req = urllib.request.Request(
            OLLAMA_URL,
            data=payload,
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=30) as resp:
            body = json.loads(resp.read())
            answer = body.get("response", "")
            # JSON aus Antwort extrahieren
            match = re.search(r"\{[^{}]+\}", answer, re.DOTALL)
            if not match:
                return {}, 0.0
            extracted = json.loads(match.group())
            return extracted, 0.75
    except (urllib.error.URLError, json.JSONDecodeError, TimeoutError):
        return {}, 0.0


def _values_to_ki(values: dict[str, Any], confidence: float) -> dict[str, dict[str, Any]]:
    """Konvertiert extrahierte Rohdaten in ki_extrahiert-Format."""
    mapping = {
        "gwp_gesamt":    "output.base.lca.gwp_gesamt",
        "pene_gesamt":   "output.base.lca.pene_gesamt",
        "bgf_m2":        "input.building.gross_floor_area",
        "nrf_m2":        "input.building.heated_area",
        "qng_plus_gwp":  "output.base.lca.qng_plus_gwp_anforderung",
        "qng_premium_gwp": "output.base.lca.qng_premium_gwp_anforderung",
    }
    result: dict[str, dict[str, Any]] = {}
    for raw_key, sidecar_pfad in mapping.items():
        val = values.get(raw_key)
        if val is not None:
            try:
                result[sidecar_pfad] = {"wert": float(val), "confidence": confidence}
            except (ValueError, TypeError):
                pass
    return result


def parse(content: bytes) -> dict[str, Any]:
    """
    Parst EVEBI Nachhaltigkeit.docx.

    Args:
        content: DOCX-Dateiinhalt als Bytes

    Returns:
        {
          "kanal": "evebi_lca_docx",
          "ki_extrahiert": { "sidecar_pfad": {"wert": X, "confidence": 0.92}, ... },
          "ki_confidence": float,   # niedrigster Einzelwert
          "warnungen": [str],
        }
    """
    if not _DOCX_AVAILABLE:
        raise RuntimeError(
            "python-docx ist nicht installiert. "
            "Bitte 'pip install python-docx' im din18599-ifc venv ausführen."
        )

    import io as _io
    doc = Document(_io.BytesIO(content))

    warnungen: list[str] = []
    ki_extrahiert: dict[str, dict[str, Any]] = {}
    confidence = 0.0

    # Schritt 1: Heuristischer Tabellen-Match
    heuristic_values: dict[str, float | None] = {}
    for table in doc.tables:
        values = _parse_qng_table(table)
        if values:
            heuristic_values.update(values)

    if heuristic_values:
        confidence = 0.92
        ki_extrahiert = _values_to_ki(heuristic_values, confidence)

        # F-Gase nicht in dieser Tabelle → Warnung für Auditor
        warnungen.append(
            "F-Gase (Kältemitteltyp + Füllmenge) konnten nicht automatisch extrahiert werden — "
            "bitte manuell im Auditor eintragen."
        )
    else:
        # Schritt 2: Ollama-Fallback
        full_text = _extract_full_text(doc)
        ollama_values, confidence = _ollama_extract(full_text)

        if ollama_values:
            ki_extrahiert = _values_to_ki(ollama_values, confidence)
            warnungen.append(
                "Tabellenstruktur nicht erkannt — Ollama-Extraktion verwendet. "
                "Auditor-Prüfung aller Werte ist Pflicht."
            )
        else:
            warnungen.append(
                "Weder Tabellen-Heuristik noch Ollama konnten Werte extrahieren. "
                "Bitte Werte manuell eintragen."
            )

    # NRF-Abweichungshinweis: Falls NRF aus DOCX von beheizter Fläche im Sidecar abweicht,
    # wird das beim Auditor-Freigabe-Schritt in der DWEapp geprüft.
    if "input.building.heated_area" in ki_extrahiert:
        warnungen.append(
            "Hinweis: NRF aus Nachhaltigkeitsbericht ≠ Wohnfläche GEG ist normal "
            "(verschiedene Normen). Abweichung wird beim Freigeben automatisch angezeigt."
        )

    return {
        "kanal": "evebi_lca_docx",
        "ki_extrahiert": ki_extrahiert,
        "ki_confidence": confidence,
        "warnungen": warnungen,
    }
